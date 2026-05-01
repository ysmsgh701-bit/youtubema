"""
YouTube 파이프라인 웹 대시보드
대본 생성 / OBS 가이드 / 영상 업로드&편집 / 파이프라인 실행 / YouTube 업로드

실행:
    python web_app.py
    브라우저: http://localhost:5000
"""
import os
import sys
import json
import uuid
import time
import builtins
import threading
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))
except ImportError:
    pass

from flask import Flask, render_template, request, jsonify, send_file, Response
import yaml

_ROOT = os.path.dirname(os.path.abspath(__file__))
CONFIG_DIR = os.path.join(_ROOT, "config")
OUTPUT_DIR = os.path.join(_ROOT, "output")
INBOX_DIR  = os.path.join(_ROOT, "inbox")

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 4 * 1024 * 1024 * 1024  # 4 GB

_jobs = {}  # job_id -> {status, logs, result, error}


def _profile(track):
    with open(os.path.join(CONFIG_DIR, f"{track}_profile.yaml"), encoding="utf-8") as f:
        return yaml.safe_load(f)


# ── Pages ──────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


# ── API: 대본 생성 ──────────────────────────────────────────────────────────

@app.route("/api/generate-script", methods=["POST"])
def api_generate_script():
    data   = request.json or {}
    track  = data.get("track", "finance")
    topic  = data.get("topic", "").strip()
    if not topic:
        return jsonify({"error": "주제를 입력하세요"}), 400
    try:
        from scripts.generate_script import generate_script, save_script
        profile  = _profile(track)
        api_key  = os.environ.get("GEMINI_API_KEY")
        script   = generate_script(topic, profile, api_key=api_key)
        path     = save_script(script, os.path.join(OUTPUT_DIR, "scripts"), topic)
        return jsonify({"script": script, "path": path})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── API: 영상 업로드 ────────────────────────────────────────────────────────

@app.route("/api/upload-video", methods=["POST"])
def api_upload_video():
    track = request.form.get("track", "finance")
    if "video" not in request.files:
        return jsonify({"error": "파일이 없습니다"}), 400
    f = request.files["video"]
    inbox = os.path.join(INBOX_DIR, track)
    os.makedirs(inbox, exist_ok=True)
    path = os.path.join(inbox, f.filename)
    f.save(path)
    return jsonify({"path": path, "filename": f.filename})


# ── API: 영상 정보 ──────────────────────────────────────────────────────────

@app.route("/api/video-info", methods=["POST"])
def api_video_info():
    path = (request.json or {}).get("video_path", "")
    if not path or not os.path.exists(path):
        return jsonify({"error": "파일 없음"}), 400
    try:
        from moviepy import VideoFileClip
        clip = VideoFileClip(path)
        info = {"duration": round(clip.duration, 2), "fps": clip.fps,
                "width": clip.size[0], "height": clip.size[1],
                "filename": os.path.basename(path), "path": path}
        clip.close()
        return jsonify(info)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── API: 영상 트림 ──────────────────────────────────────────────────────────

@app.route("/api/trim-video", methods=["POST"])
def api_trim_video():
    data  = request.json or {}
    path  = data.get("video_path", "")
    start = float(data.get("start", 0))
    end   = float(data.get("end", 0))
    if not path or not os.path.exists(path):
        return jsonify({"error": "파일 없음"}), 400
    try:
        from moviepy import VideoFileClip
        clip = VideoFileClip(path)
        if end <= 0 or end > clip.duration:
            end = clip.duration
        if start >= end:
            clip.close()
            return jsonify({"error": "start >= end"}), 400
        base, ext = os.path.splitext(path)
        out = f"{base}_trimmed{ext}"
        clip.subclipped(start, end).write_videofile(out, logger=None)
        clip.close()
        return jsonify({"path": out, "duration": round(end - start, 2)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── API: 파이프라인 실행 ────────────────────────────────────────────────────

@app.route("/api/run-pipeline", methods=["POST"])
def api_run_pipeline():
    data       = request.json or {}
    video_path = data.get("video_path", "")
    track      = data.get("track", "finance")
    dry_run    = data.get("dry_run", True)
    model      = data.get("model", "small")
    privacy    = data.get("privacy", "private")

    if not video_path or not os.path.exists(video_path):
        return jsonify({"error": f"영상 없음: {video_path}"}), 400

    job_id = str(uuid.uuid4())[:8]
    _jobs[job_id] = {"status": "running", "logs": [], "result": None, "error": None}

    threading.Thread(
        target=_bg_pipeline,
        args=(job_id, video_path, track, dry_run, model, privacy),
        daemon=True,
    ).start()

    return jsonify({"job_id": job_id})


def _bg_pipeline(job_id, video_path, track, dry_run, model, privacy):
    orig = builtins.print
    def capture(*args, **kw):
        msg = " ".join(str(a) for a in args)
        if msg.strip():
            _jobs[job_id]["logs"].append(msg.strip())
        orig(*args, **kw)
    builtins.print = capture
    try:
        from pipeline import run_pipeline
        result = run_pipeline(video_path, track, dry_run=dry_run,
                              whisper_model=model, privacy=privacy)
        _jobs[job_id]["status"] = "done"
        _jobs[job_id]["result"] = result
    except Exception as e:
        _jobs[job_id]["status"] = "error"
        _jobs[job_id]["error"]  = str(e)
        _jobs[job_id]["logs"].append(f"[ERROR] {e}")
    finally:
        builtins.print = orig


# ── API: 파이프라인 진행 (SSE) ──────────────────────────────────────────────

@app.route("/api/progress/<job_id>")
def api_progress(job_id):
    def stream():
        sent = 0
        while True:
            if job_id not in _jobs:
                yield f"data: {json.dumps({'error': 'unknown job'})}\n\n"
                return
            job  = _jobs[job_id]
            logs = job["logs"]
            while sent < len(logs):
                yield f"data: {json.dumps({'log': logs[sent]})}\n\n"
                sent += 1
            if job["status"] == "done":
                yield f"data: {json.dumps({'done': True, 'result': job['result']})}\n\n"
                return
            if job["status"] == "error":
                yield f"data: {json.dumps({'error': job['error']})}\n\n"
                return
            time.sleep(0.4)
    return Response(stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── API: DOCX 내보내기 ──────────────────────────────────────────────────────

@app.route("/api/export-docx", methods=["POST"])
def api_export_docx():
    data    = request.json or {}
    script  = data.get("script", "").strip()
    topic   = data.get("topic", "script")
    track   = data.get("track", "finance")
    if not script:
        return jsonify({"error": "대본 내용이 없습니다"}), 400
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 스타일 기본 설정
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"
        style.font.size = Pt(11)

        # 제목
        title = doc.add_heading(topic, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xDB)

        doc.add_paragraph(f"트랙: {track} | 생성일: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph("")

        # 마크다운 → Word 변환 (헤딩/본문)
        for line in script.split("\n"):
            if line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        safe = "".join(c for c in topic if c.isalnum() or c in " _-")[:30].strip()
        filename = f"{safe}_script.docx"
        out_path = os.path.join(OUTPUT_DIR, "scripts", filename)
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)

        return send_file(out_path, as_attachment=True,
                         download_name=filename,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── API: main_macro 파이프라인 실행 (AI 영상) ───────────────────────────────

@app.route("/api/run-ai-video", methods=["POST"])
def api_run_ai_video():
    data    = request.json or {}
    topic   = data.get("topic", "").strip()
    langs   = data.get("langs", "ko")
    formats = data.get("formats", "longform")
    dry_run = data.get("dry_run", True)
    pick    = data.get("pick", None)
    if not topic:
        return jsonify({"error": "주제를 입력하세요"}), 400

    job_id = str(uuid.uuid4())[:8]
    _jobs[job_id] = {"status": "running", "logs": [], "result": None, "error": None}

    threading.Thread(
        target=_bg_ai_video,
        args=(job_id, topic, langs, formats, dry_run, pick),
        daemon=True,
    ).start()
    return jsonify({"job_id": job_id})


def _bg_ai_video(job_id, topic, langs, formats, dry_run, pick):
    orig = builtins.print
    def capture(*args, **kw):
        msg = " ".join(str(a) for a in args)
        if msg.strip():
            _jobs[job_id]["logs"].append(msg.strip())
        orig(*args, **kw)
    builtins.print = capture
    try:
        from main_macro import YouTubeAutomationPipelineV2
        api_key = os.environ.get("GEMINI_API_KEY")
        pipeline = YouTubeAutomationPipelineV2(
            languages=[l.strip() for l in langs.split(",")],
            formats=[f.strip() for f in formats.split(",")],
            dry_run=dry_run,
            api_key=api_key,
        )
        pipeline.run_phases(start_phase=1, end_phase=6, auto_pick=pick)
        _jobs[job_id]["status"] = "done"
        _jobs[job_id]["result"] = {"run_dir": pipeline.run_dir}
    except Exception as e:
        _jobs[job_id]["status"] = "error"
        _jobs[job_id]["error"]  = str(e)
        _jobs[job_id]["logs"].append(f"[ERROR] {e}")
    finally:
        builtins.print = orig


# ── API: 결과 목록 ──────────────────────────────────────────────────────────

@app.route("/api/outputs/<track>")
def api_outputs(track):
    track_dir = os.path.join(OUTPUT_DIR, track)
    if not os.path.isdir(track_dir):
        return jsonify([])
    items = []
    for name in sorted(os.listdir(track_dir), reverse=True)[:15]:
        run_dir = os.path.join(track_dir, name)
        if not os.path.isdir(run_dir):
            continue
        item = {"timestamp": name, "run_dir": run_dir}
        mp = os.path.join(run_dir, "meta.json")
        if os.path.exists(mp):
            with open(mp, encoding="utf-8") as f:
                item["meta"] = json.load(f)
        item["has_thumbnail"] = os.path.exists(os.path.join(run_dir, "thumbnail.jpg"))
        item["has_srt"]       = os.path.exists(os.path.join(run_dir, "subtitles.srt"))
        rp = os.path.join(run_dir, "pipeline_result.json")
        if os.path.exists(rp):
            with open(rp, encoding="utf-8") as f:
                item["result"] = json.load(f)
        items.append(item)
    return jsonify(items)


# ── API: 파일 서빙 ──────────────────────────────────────────────────────────

@app.route("/api/file")
def api_file():
    path = request.args.get("path", "")
    if not path:
        return "missing path", 400
    abs_path = os.path.abspath(path)
    if not abs_path.startswith(_ROOT) or not os.path.exists(abs_path):
        return "not found", 404
    return send_file(abs_path)


@app.route("/api/open-folder", methods=["POST"])
def api_open_folder():
    track = (request.json or {}).get("track", "finance")
    folder = os.path.join(INBOX_DIR, track)
    os.makedirs(folder, exist_ok=True)
    try:
        import subprocess
        subprocess.Popen(["explorer", folder])
    except Exception:
        pass
    return jsonify({"folder": folder})


@app.route("/api/srt-text")
def api_srt_text():
    path = request.args.get("path", "")
    abs_path = os.path.abspath(path)
    if not abs_path.startswith(_ROOT) or not os.path.exists(abs_path):
        return "not found", 404
    with open(abs_path, encoding="utf-8") as f:
        return f.read(), 200, {"Content-Type": "text/plain; charset=utf-8"}


if __name__ == "__main__":
    print("=" * 52)
    print("  YouTube 파이프라인 웹 대시보드")
    print("  http://localhost:5001")
    print("=" * 52)
    app.run(debug=False, host="0.0.0.0", port=5001, threaded=True)
